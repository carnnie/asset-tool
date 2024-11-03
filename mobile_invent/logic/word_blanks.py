from docx import Document
from docx.text.paragraph import Paragraph

import logging

logger = logging.getLogger("main")

class WordDocument:
    def __init__(self, path: str) -> None:
        self.document = Document(path)
        self.data: dict[str, str] = {}

    def set_data(self, data: dict) -> bool:
        logger.warning(data)
        self.data = {**data} if data else self.data
        return bool(self.data)

    def save(self, path) -> None:
        self.document.save(path)

    def to_change(self, paragraph: Paragraph) -> tuple[str, str] | None:
        for word in self.data.keys():
            if word in paragraph.text:
                return (word, self.data.pop(word))
        return None

    def change_table_text(self, data: dict = {}) -> None:
        if not self.set_data(data):
            return None
        """sorry for indentation=)"""
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if pair := self.to_change(paragraph):
                            paragraph.text = paragraph.text.replace(*pair)
                            if not self.data:
                                return None

    def change_paragraph_text(self, data) -> None:
        if not self.set_data(data):
            return None

    def change_document_text(self, data: dict) -> Document:
        if not self.set_data(data):
            return None
        self.change_table_text(data)
        self.change_paragraph_text(data)
        return self.document


if __name__ == "__main__":
    doc = WordDocument("./static/blanks/empty/giveaway/laptop.docx")
    doc.change_table_text({"$DEVICE": "IPAD", "$NUMBER": "123"})
    doc.save("./test.docx")
